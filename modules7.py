# Дополнительно устанавливаемые модули
# Модули документов
# PIP - Python Installation Package
# pip install python-docx
# создание и наполнение документа Word
# https://python-docx.readthedocs.io/en/latest/
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

document = Document()  # создали бланк документа

header = document.add_heading('Заголовок', 0)
header.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_paragraph('Абзац с форматированием по умолчанию.')

# Немного сложный абзац
paragraph = document.add_paragraph('Обычный текст, ')
paragraph.add_run('а это жирный текст, ').bold = True
paragraph.add_run('а это наклонный текст.').italic = True

document.add_paragraph('Элемент unordered list',
                       style='List Bullet')

document.add_paragraph('Элемент ordered list',
                       style='List Number')

document.add_picture('images/flag.jpg', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.save('word/test.docx')
